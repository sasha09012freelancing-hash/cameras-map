import xml.etree.ElementTree as ET
import re, json, math
from docx import Document
import openpyxl, xlrd
import pdfplumber

OUTPUT_FILE = "cameras_map.html"

REGIONS = {
    "rostov":        {"name": "Ростовская обл.",        "color": "#e53935", "file": "cameras_rostov.kml"},
    "adygea":        {"name": "Республика Адыгея",      "color": "#fb8c00", "file": "adygea_cameras.docx"},
    "crimea":        {"name": "Республика Крым",        "color": "#8e24aa", "file": "crimea_cameras.xlsx"},
    "krasnodar":     {"name": "Краснодарский край",     "color": "#00897b", "file": "krasnodar_cameras.pdf"},
    "krasnodar_osm": {"name": "Краснодарский край",     "color": "#00897b", "file": "krasnodar_osm.json"},
    "stavropol":     {"name": "Ставропольский край",    "color": "#3949ab", "file": "stavropol.xlsx"},
    "dagestan":      {"name": "Республика Дагестан",    "color": "#c0392b", "file": "dagestan.xlsx"},
    "kalmykia":      {"name": "Республика Калмыкия",    "color": "#7f8c8d", "file": "kalmykia.xlsx"},
    "astrakhan":     {"name": "Астраханская обл.",      "color": "#d35400", "file": "astrakhan.docx"},
    "volgograd":     {"name": "Волгоградская обл.",     "color": "#27ae60", "file": "volgograd.xls"},
    "kbr":           {"name": "Кабардино-Балкария",     "color": "#16a085", "file": "kbr.xlsx"},
    "north_ossetia": {"name": "Сев. Осетия",            "color": "#8e44ad", "file": "north_ossetia.xlsx"},
    "kchr":          {"name": "Карачаево-Черкесия",     "color": "#2980b9", "file": "kchr.xlsx"},
    "chechnya":      {"name": "Чеченская Республика",   "color": "#e67e22", "file": "chechnya.xlsx"},
    "ingushetia":    {"name": "Республика Ингушетия",   "color": "#1abc9c", "file": "ingushetia.xlsx"},
}

def is_valid(lat, lon):
    try:
        return (not math.isnan(lat) and not math.isnan(lon)
                and -90 <= lat <= 90 and -180 <= lon <= 180)
    except:
        return False

def detect_type(violations):
    v = (violations or "").lower()
    if any(w in v for w in ["паркова", "стоянк", "остановк"]):
        return "parking"
    return "speed"

def guess_speed(address, violations=""):
    a = (address + " " + violations).lower()
    if re.search(r'\bм-\d|\bа-\d|\bр-\d|\bе-\d|а/д\b|автодор|трасс|\d+\s*км\b', a):
        # Federal/regional highway → 90 km/h
        if re.search(r'м-4|м-25|м-23|е-50|трасс', a):
            return "90"
        return "90"
    if re.search(r'ул\.|пр\.|пр-кт|пер\.|наб\.|бульв|алл\.|мкр|микрорай', a):
        return "60"
    if re.search(r'\bг\.\s|\bгород\b|\bст\.\s|\bпос\.\s', a):
        return "60"
    return ""

def to_float(s):
    try:
        return float(str(s).replace(",", ".").strip())
    except:
        return float("nan")

# --- Ростов (KML) ---
def parse_kml(filepath):
    cameras = []
    ns = {"kml": "http://www.opengis.net/kml/2.2"}
    tree = ET.parse(filepath)
    doc = tree.getroot().find("kml:Document", ns)
    for pm in doc.findall("kml:Placemark", ns):
        coords_el = pm.find(".//kml:coordinates", ns)
        if coords_el is None:
            continue
        parts = coords_el.text.strip().split(",")
        if len(parts) < 2:
            continue
        lon_f, lat_f = to_float(parts[0]), to_float(parts[1])
        if not is_valid(lat_f, lon_f):
            continue
        desc = (pm.find("kml:description", ns) or pm).text or ""
        addr  = (re.search(r"<b>Адрес:</b>\s*(.*?)<br", desc, re.DOTALL) or type('',(),{'group':lambda s,i:''})()).group(1).strip()
        viol  = (re.search(r"<b>Нарушения:</b>\s*(.*?)<br", desc, re.DOTALL) or type('',(),{'group':lambda s,i:''})()).group(1).strip()
        cond  = (re.search(r"<b>Условия:</b>\s*(.*?)<br", desc, re.DOTALL) or type('',(),{'group':lambda s,i:''})()).group(1).strip()
        name_el = pm.find("kml:name", ns)
        style_el = pm.find("kml:styleUrl", ns)
        cam_type = style_el.text.strip("#") if style_el is not None else "speed"
        cameras.append({"lat": lat_f, "lon": lon_f,
                        "name": name_el.text.strip() if name_el is not None else addr,
                        "address": addr, "violations": viol, "conditions": cond,
                        "speed": guess_speed(addr, viol),
                        "type": cam_type, "region": "rostov"})
    return cameras

# --- Адыгея (DOCX) ---
def parse_docx(filepath):
    cameras = []
    doc = Document(filepath)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 3:
                continue
            lat_f = to_float(cells[1])
            lon_f = to_float(cells[2])
            if not is_valid(lat_f, lon_f):
                continue
            addr  = cells[0]
            viol  = cells[3] if len(cells) > 3 else ""
            cond  = cells[4] if len(cells) > 4 else ""
            cameras.append({"lat": lat_f, "lon": lon_f,
                            "name": addr, "address": addr,
                            "violations": viol, "conditions": cond,
                            "speed": guess_speed(addr, viol),
                            "type": detect_type(viol), "region": "adygea"})
    return cameras

# --- Универсальный XLSX парсер ---
# col_addr, col_lat, col_lon, col_viol, col_cond — индексы столбцов (0-based)
def parse_xlsx_generic(filepath, region, skip=3, col_addr=0, col_lat=1, col_lon=2, col_viol=3, col_cond=4):
    cameras = []
    wb = openpyxl.load_workbook(filepath)
    ws = wb.active
    for row in ws.iter_rows(min_row=skip+1, values_only=True):
        if not row: continue
        lat_f = to_float(row[col_lat] if col_lat < len(row) else "")
        lon_f = to_float(row[col_lon] if col_lon < len(row) else "")
        if not is_valid(lat_f, lon_f): continue
        addr = str(row[col_addr] or "") if col_addr < len(row) else ""
        viol = str(row[col_viol] or "") if col_viol < len(row) else ""
        cond = str(row[col_cond] or "") if col_cond < len(row) else ""
        cameras.append({"lat": lat_f, "lon": lon_f,
                        "name": addr, "address": addr,
                        "violations": viol, "conditions": cond,
                        "speed": guess_speed(addr, viol),
                        "type": detect_type(viol), "region": region})
    return cameras

# --- Крым (XLSX) ---
def parse_xlsx(filepath):
    return parse_xlsx_generic(filepath, "crimea", skip=3, col_addr=0, col_lat=1, col_lon=2, col_viol=3, col_cond=4)

# --- Волгоград (старый XLS) ---
def parse_xls(filepath, region):
    cameras = []
    try:
        wb = xlrd.open_workbook(filepath)
        ws = wb.sheet_by_index(0)
        for i in range(3, ws.nrows):
            try:
                addr = str(ws.cell_value(i, 0) or "")
                lat_f = to_float(ws.cell_value(i, 1))
                lon_f = to_float(ws.cell_value(i, 2))
                if not is_valid(lat_f, lon_f): continue
                viol = str(ws.cell_value(i, 3) or "")
                cond = str(ws.cell_value(i, 4) or "") if ws.ncols > 4 else ""
                cameras.append({"lat": lat_f, "lon": lon_f,
                                "name": addr, "address": addr,
                                "violations": viol, "conditions": cond,
                                "speed": guess_speed(addr, viol),
                                "type": detect_type(viol), "region": region})
            except: continue
    except Exception as e:
        print(f"  XLS ERR {filepath}: {e}")
    return cameras

# --- Калмыкия (координаты в тексте) ---
def parse_kalmykia(filepath):
    cameras = []
    wb = openpyxl.load_workbook(filepath)
    ws = wb.active
    for row in ws.iter_rows(min_row=2, values_only=True):
        if not row or not row[0]: continue
        cell = str(row[0])
        viol = str(row[1] or "") if len(row) > 1 else ""
        # Parse "lat, lon; lat, lon; ..." pairs
        pairs = re.findall(r'(\d{2,3}\.\d+)[,\s]+(\d{2,3}\.\d+)', cell)
        for lat_s, lon_s in pairs:
            lat_f, lon_f = to_float(lat_s), to_float(lon_s)
            # Russian coords: lat ~44-48, lon ~40-50
            if not is_valid(lat_f, lon_f): continue
            if not (40 <= lat_f <= 60 and 30 <= lon_f <= 60): continue
            cameras.append({"lat": lat_f, "lon": lon_f,
                            "name": cell[:60], "address": cell[:60],
                            "violations": viol, "conditions": "",
                            "speed": guess_speed(cell, viol),
                            "type": detect_type(viol), "region": "kalmykia"})
    return cameras

# --- Краснодар OSM (JSON) ---
SPEED_MAP = {"RU:urban": "60", "RU:rural": "90", "RU:motorway": "110"}

def parse_osm_json(filepath):
    cameras = []
    with open(filepath, encoding="utf-8") as f:
        nodes = json.load(f)
    for node in nodes:
        lat_f = to_float(node.get("lat", "nan"))
        lon_f = to_float(node.get("lon", "nan"))
        if not is_valid(lat_f, lon_f):
            continue
        tags = node.get("tags", {})
        addr  = tags.get("name", tags.get("description", tags.get("addr:street", "")))
        raw_speed = tags.get("maxspeed", "")
        speed = SPEED_MAP.get(raw_speed, raw_speed)
        osm_id = node.get("id", "")
        cameras.append({"lat": lat_f, "lon": lon_f,
                        "name": addr or f"{lat_f:.4f},{lon_f:.4f}",
                        "address": addr, "violations": "Превышение скорости",
                        "conditions": tags.get("direction", ""),
                        "speed": speed, "osm_id": osm_id,
                        "type": "speed", "region": "krasnodar_osm"})
    return cameras

# --- Краснодар (PDF) ---
def parse_pdf(filepath):
    cameras = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            for table in page.extract_tables():
                for row in table:
                    if not row or not row[1]:
                        continue
                    lat_f = to_float(row[1])
                    lon_f = to_float(row[2])
                    if not is_valid(lat_f, lon_f):
                        continue
                    addr  = str(row[0] or "").replace("\n", " ")
                    viol  = str(row[3] or "").replace("\n", " ")
                    cond  = str(row[4] or "").replace("\n", " ")
                    cameras.append({"lat": lat_f, "lon": lon_f,
                                    "name": addr, "address": addr,
                                    "violations": viol, "conditions": cond,
                                    "speed": guess_speed(addr, viol),
                                    "type": detect_type(viol), "region": "krasnodar"})
    return cameras

# --- Астрахань (DOCX таблица) ---
def parse_docx_generic(filepath, region, skip=3, col_addr=1, col_lat=2, col_lon=3):
    cameras = []
    doc = Document(filepath)
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            if i < skip:
                continue
            cells = [c.text.strip() for c in row.cells]
            if len(cells) <= max(col_addr, col_lat, col_lon):
                continue
            lat_f = to_float(cells[col_lat])
            lon_f = to_float(cells[col_lon])
            if not is_valid(lat_f, lon_f):
                continue
            addr = cells[col_addr]
            viol = cells[4] if len(cells) > 4 else ""
            cameras.append({"lat": lat_f, "lon": lon_f,
                            "name": addr, "address": addr,
                            "violations": viol, "conditions": "",
                            "speed": guess_speed(addr, viol),
                            "type": detect_type(viol), "region": region})
    return cameras

# --- Собираем всё ---
all_cameras = []
parsers = {
    "rostov":        (parse_kml,      REGIONS["rostov"]["file"]),
    "adygea":        (parse_docx,     REGIONS["adygea"]["file"]),
    "crimea":        (parse_xlsx,     REGIONS["crimea"]["file"]),
    "krasnodar":     (parse_pdf,      REGIONS["krasnodar"]["file"]),
    "krasnodar_osm": (parse_osm_json, REGIONS["krasnodar_osm"]["file"]),
    "stavropol":     (lambda f: parse_xlsx_generic(f, "stavropol", skip=4, col_addr=2, col_lat=4, col_lon=5, col_viol=6, col_cond=7), REGIONS["stavropol"]["file"]),
    "dagestan":      (lambda f: parse_xlsx_generic(f, "dagestan",  skip=3, col_addr=1, col_lat=2, col_lon=3, col_viol=4, col_cond=5), REGIONS["dagestan"]["file"]),
    "kalmykia":      (parse_kalmykia, REGIONS["kalmykia"]["file"]),
    "astrakhan":     (lambda f: parse_docx_generic(f, "astrakhan"), REGIONS["astrakhan"]["file"]),
    "volgograd":     (lambda f: parse_xls(f, "volgograd"),          REGIONS["volgograd"]["file"]),
    "kbr":           (lambda f: parse_xlsx_generic(f, "kbr"),       REGIONS["kbr"]["file"]),
    "north_ossetia": (lambda f: parse_xlsx_generic(f, "north_ossetia"), REGIONS["north_ossetia"]["file"]),
    "kchr":          (lambda f: parse_xlsx_generic(f, "kchr"),      REGIONS["kchr"]["file"]),
    "chechnya":      (lambda f: parse_xlsx_generic(f, "chechnya",   skip=0, col_addr=1, col_lat=2, col_lon=3, col_viol=4, col_cond=5), REGIONS["chechnya"]["file"]),
    "ingushetia":    (lambda f: parse_xlsx_generic(f, "ingushetia"), REGIONS["ingushetia"]["file"]),
}

stats = {}
for key, (fn, fpath) in parsers.items():
    try:
        cams = fn(fpath)
        stats[key] = len(cams)
        all_cameras.extend(cams)
        print(f"  {REGIONS[key]['name']}: {len(cams)} камер")
    except Exception as e:
        print(f"  ОШИБКА {key}: {e}")
        stats[key] = 0

print(f"\nИтого: {len(all_cameras)} камер")

cameras_json = json.dumps(all_cameras, ensure_ascii=False)
regions_json = json.dumps({k: {"name": v["name"], "color": v["color"]} for k, v in REGIONS.items()}, ensure_ascii=False)

html = f"""<!DOCTYPE html>
<html lang="ru">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=1.0">
  <title>Камеры — Юг России</title>
  <link rel="stylesheet" href="https://unpkg.com/leaflet@1.9.4/dist/leaflet.css"/>
  <style>
    * {{ margin:0; padding:0; box-sizing:border-box; }}
    html,body {{ height:100%; font-family:Arial,sans-serif; background:#1a1a2e; }}
    #map {{ width:100%; height:100vh; }}

    #panel {{
      position:fixed; top:10px; left:10px; z-index:1000;
      background:rgba(22,26,40,0.95); border-radius:12px;
      padding:10px 14px; box-shadow:0 4px 20px rgba(0,0,0,0.5);
      min-width:205px; max-width:240px;
      max-height:calc(100vh - 20px); overflow-y:auto;
      border:1px solid rgba(255,255,255,0.08);
    }}
    #panel h3 {{ font-size:13px; color:#e0e0e0; margin-bottom:8px; border-bottom:1px solid rgba(255,255,255,0.1); padding-bottom:6px; }}
    #panel label {{
      display:flex; align-items:center; gap:7px;
      font-size:12px; cursor:pointer; margin:5px 0; user-select:none; color:#ccc;
    }}
    #panel label:hover {{ color:#fff; }}
    #panel input[type=checkbox] {{ width:14px; height:14px; cursor:pointer; accent-color:#4fc3f7; }}
    .dot {{ width:11px; height:11px; border-radius:50%; flex-shrink:0; border:1.5px solid rgba(255,255,255,0.25); }}
    .badge {{ font-size:10px; color:#666; margin-left:auto; }}
    .sep {{ border-top:1px solid rgba(255,255,255,0.08); margin:7px 0; }}

    /* Tile switcher */
    #tilebtn {{
      position:fixed; top:10px; right:54px; z-index:1000;
      background:rgba(22,26,40,0.95); border:1px solid rgba(255,255,255,0.1);
      border-radius:10px; padding:6px 8px;
      box-shadow:0 4px 15px rgba(0,0,0,0.4);
      display:flex; gap:5px;
    }}
    #tilebtn button {{
      border:none; border-radius:7px; padding:5px 10px;
      font-size:11px; cursor:pointer; font-weight:600;
      transition:all 0.2s;
    }}
    #tilebtn button.active {{ background:#4fc3f7; color:#000; }}
    #tilebtn button:not(.active) {{ background:rgba(255,255,255,0.1); color:#aaa; }}
    #tilebtn button:hover:not(.active) {{ background:rgba(255,255,255,0.2); color:#fff; }}

    #alert {{
      position:fixed; bottom:20px; left:50%; transform:translateX(-50%);
      z-index:2000; background:#e53935; color:#fff;
      padding:12px 22px; border-radius:10px;
      font-size:15px; font-weight:bold;
      box-shadow:0 3px 12px rgba(0,0,0,0.5);
      display:none; text-align:center; max-width:90vw;
    }}
    #locbtn {{
      position:fixed; bottom:20px; right:10px; z-index:1000;
      background:rgba(22,26,40,0.95); border:1px solid rgba(255,255,255,0.1);
      border-radius:50%; width:44px; height:44px; cursor:pointer;
      box-shadow:0 2px 10px rgba(0,0,0,0.4); font-size:20px;
    }}
    .popup-title {{ font-weight:bold; font-size:13px; margin-bottom:5px; color:#111; }}
    .popup-row {{ font-size:12px; line-height:1.6; color:#333; }}
    .popup-row b {{ color:#111; }}
    .popup-region {{ font-size:11px; color:#888; margin-top:4px; }}
    .speed-badge {{
      display:inline-block; background:#e53935; color:#fff;
      font-size:15px; font-weight:bold; border-radius:6px;
      padding:2px 10px; margin-bottom:6px;
    }}
    /* Leaflet dark popup */
    .leaflet-popup-content-wrapper {{ border-radius:10px; }}
  </style>
</head>
<body>
  <div id="map"></div>

  <div id="tilebtn">
    <button id="tile-dark" class="active" onclick="setTile('dark')">🌙 Тёмная</button>
    <button id="tile-street" onclick="setTile('street')">🗺 Дороги</button>
    <button id="tile-light" onclick="setTile('light')">☀️ Светлая</button>
  </div>

  <div id="panel">
    <h3>📷 Камеры — Юг России</h3>
    <label><input type="checkbox" id="chk-rostov" checked>
      <span class="dot" style="background:#e53935"></span>Ростовская обл.
      <span class="badge">{stats.get('rostov',0)}</span></label>
    <label><input type="checkbox" id="chk-krasnodar" checked>
      <span class="dot" style="background:#00897b"></span>Краснодарский край
      <span class="badge">{stats.get('krasnodar',0) + stats.get('krasnodar_osm',0)}</span></label>
    <label><input type="checkbox" id="chk-crimea" checked>
      <span class="dot" style="background:#8e24aa"></span>Крым
      <span class="badge">{stats.get('crimea',0)}</span></label>
    <label><input type="checkbox" id="chk-adygea" checked>
      <span class="dot" style="background:#fb8c00"></span>Адыгея
      <span class="badge">{stats.get('adygea',0)}</span></label>
    <label><input type="checkbox" id="chk-stavropol" checked>
      <span class="dot" style="background:#3949ab"></span>Ставропольский край
      <span class="badge">{stats.get('stavropol',0)}</span></label>
    <label><input type="checkbox" id="chk-volgograd" checked>
      <span class="dot" style="background:#27ae60"></span>Волгоградская обл.
      <span class="badge">{stats.get('volgograd',0)}</span></label>
    <label><input type="checkbox" id="chk-astrakhan" checked>
      <span class="dot" style="background:#d35400"></span>Астраханская обл.
      <span class="badge">{stats.get('astrakhan',0)}</span></label>
    <label><input type="checkbox" id="chk-dagestan" checked>
      <span class="dot" style="background:#c0392b"></span>Дагестан
      <span class="badge">{stats.get('dagestan',0)}</span></label>
    <label><input type="checkbox" id="chk-kalmykia" checked>
      <span class="dot" style="background:#7f8c8d"></span>Калмыкия
      <span class="badge">{stats.get('kalmykia',0)}</span></label>
    <label><input type="checkbox" id="chk-kbr" checked>
      <span class="dot" style="background:#16a085"></span>Кабардино-Балкария
      <span class="badge">{stats.get('kbr',0)}</span></label>
    <label><input type="checkbox" id="chk-north_ossetia" checked>
      <span class="dot" style="background:#8e44ad"></span>Сев. Осетия
      <span class="badge">{stats.get('north_ossetia',0)}</span></label>
    <label><input type="checkbox" id="chk-kchr" checked>
      <span class="dot" style="background:#2980b9"></span>Карачаево-Черкесия
      <span class="badge">{stats.get('kchr',0)}</span></label>
    <label><input type="checkbox" id="chk-chechnya" checked>
      <span class="dot" style="background:#e67e22"></span>Чечня
      <span class="badge">{stats.get('chechnya',0)}</span></label>
    <label><input type="checkbox" id="chk-ingushetia" checked>
      <span class="dot" style="background:#1abc9c"></span>Ингушетия
      <span class="badge">{stats.get('ingushetia',0)}</span></label>
    <div class="sep"></div>
    <label><input type="checkbox" id="chk-warn" checked>⚠️ Предупреждения (300м)</label>
  </div>
  <div id="alert"></div>
  <button id="locbtn" title="Моё местоположение">📍</button>

  <script src="https://unpkg.com/leaflet@1.9.4/dist/leaflet.js"></script>
  <script>
    const CAMERAS = {cameras_json};
    const REGIONS = {regions_json};
    const WARN_RADIUS = 300;

    const TILES = {{
      dark:   {{ url:"https://{{s}}.basemaps.cartocdn.com/dark_all/{{z}}/{{x}}/{{y}}.png",   attr:"&copy; OpenStreetMap &copy; CARTO" }},
      street: {{ url:"https://server.arcgisonline.com/ArcGIS/rest/services/World_Street_Map/MapServer/tile/{{z}}/{{y}}/{{x}}", attr:"&copy; Esri" }},
      light:  {{ url:"https://{{s}}.basemaps.cartocdn.com/rastertiles/voyager/{{z}}/{{x}}/{{y}}.png", attr:"&copy; OpenStreetMap &copy; CARTO" }},
    }};

    const map = L.map("map", {{zoomControl:false}}).setView([46.5, 38.5], 7);
    L.control.zoom({{position:"topright"}}).addTo(map);

    let currentTile = L.tileLayer(TILES.dark.url, {{
      attribution: TILES.dark.attr, subdomains:"abcd", maxZoom:19
    }}).addTo(map);
    let currentTileKey = "dark";

    window.setTile = function(key) {{
      map.removeLayer(currentTile);
      currentTile = L.tileLayer(TILES[key].url, {{
        attribution: TILES[key].attr, subdomains:"abcd", maxZoom:19
      }}).addTo(map);
      currentTile.bringToBack();
      currentTileKey = key;
      document.querySelectorAll("#tilebtn button").forEach(b => b.classList.remove("active"));
      document.getElementById("tile-"+key).classList.add("active");
    }};

    const layers = {{}};
    Object.keys(REGIONS).forEach(r => {{ layers[r] = L.layerGroup().addTo(map); }});

    CAMERAS.forEach(cam => {{
      const color = REGIONS[cam.region] ? REGIONS[cam.region].color : "#999";
      const marker = L.circleMarker([cam.lat, cam.lon], {{
        radius:7, fillColor:color, color:"rgba(0,0,0,0.4)",
        weight:1.5, opacity:1, fillOpacity:0.92
      }});
      const regionName = REGIONS[cam.region] ? REGIONS[cam.region].name : cam.region;
      const speedBadge = cam.speed ? `<div><span class="speed-badge">${{cam.speed}} км/ч</span></div>` : '';
      const osmLink = cam.osm_id ? `<div class="popup-region"><a href="https://www.openstreetmap.org/node/${{cam.osm_id}}" target="_blank">OpenStreetMap</a></div>` : '';
      const addrRow = cam.address ? `<div class="popup-row"><b>Адрес:</b> ${{cam.address}}</div>` : '';
      marker.bindPopup(`
        ${{speedBadge}}
        <div class="popup-title">${{cam.name}}</div>
        ${{addrRow}}
        <div class="popup-row"><b>Нарушения:</b> ${{cam.violations}}</div>
        <div class="popup-row"><b>Условия:</b> ${{cam.conditions}}</div>
        <div class="popup-region">${{regionName}}</div>
        ${{osmLink}}
      `, {{maxWidth:300}});
      marker.camData = cam;
      if (layers[cam.region]) layers[cam.region].addLayer(marker);
    }});

    ["rostov","adygea","crimea","stavropol","dagestan","kalmykia",
     "astrakhan","volgograd","kbr","north_ossetia","kchr","chechnya","ingushetia"].forEach(r => {{
      const el = document.getElementById("chk-" + r);
      if (el) el.addEventListener("change", e => {{
        e.target.checked ? layers[r].addTo(map) : map.removeLayer(layers[r]);
      }});
    }});
    document.getElementById("chk-krasnodar").addEventListener("change", e => {{
      ["krasnodar","krasnodar_osm"].forEach(r => {{
        e.target.checked ? layers[r].addTo(map) : map.removeLayer(layers[r]);
      }});
    }});

    let userMarker=null, watchId=null, lastAlert=0, alertTO=null;
    const alertEl = document.getElementById("alert");

    function showAlert(txt) {{
      alertEl.innerHTML = txt;
      alertEl.style.display = "block";
      clearTimeout(alertTO);
      alertTO = setTimeout(() => {{ alertEl.style.display="none"; }}, 5000);
    }}

    function checkNear(lat, lon) {{
      if (!document.getElementById("chk-warn").checked) return;
      const now = Date.now();
      if (now - lastAlert < 8000) return;
      let nearest=null, minD=Infinity;
      CAMERAS.forEach(cam => {{
        const d = map.distance([lat,lon],[cam.lat,cam.lon]);
        if (d < WARN_RADIUS && d < minD) {{ minD=d; nearest=cam; }}
      }});
      if (nearest) {{
        lastAlert = now;
        showAlert(`⚠️ Камера через ${{Math.round(minD)}} м<br><small>${{nearest.violations.slice(0,60)}}</small>`);
      }}
    }}

    function onPos(pos) {{
      const lat=pos.coords.latitude, lon=pos.coords.longitude;
      if (!userMarker) {{
        userMarker = L.circleMarker([lat,lon],{{
          radius:8, fillColor:"#29b6f6", color:"#fff", weight:3, fillOpacity:1
        }}).addTo(map);
      }} else {{ userMarker.setLatLng([lat,lon]); }}
      checkNear(lat,lon);
    }}

    document.getElementById("locbtn").addEventListener("click", () => {{
      if (!navigator.geolocation) {{ alert("Геолокация не поддерживается"); return; }}
      if (watchId !== null) {{
        navigator.geolocation.clearWatch(watchId); watchId=null;
        if(userMarker){{ map.removeLayer(userMarker); userMarker=null; }}
        document.getElementById("locbtn").style.background="rgba(22,26,40,0.95)"; return;
      }}
      document.getElementById("locbtn").style.background="#1a3a5c";
      navigator.geolocation.getCurrentPosition(p => {{
        map.setView([p.coords.latitude,p.coords.longitude],13); onPos(p);
      }}, err => alert("Ошибка геолокации: "+err.message));
      watchId = navigator.geolocation.watchPosition(onPos, null,
        {{enableHighAccuracy:true, maximumAge:3000, timeout:10000}});
    }});
  </script>
</body>
</html>"""

with open(OUTPUT_FILE, "w", encoding="utf-8") as f:
    f.write(html)
print(f"\nГотово! Файл: {OUTPUT_FILE}")
